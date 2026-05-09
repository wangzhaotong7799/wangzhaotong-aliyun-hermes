#!/root/.hermes/hermes-agent/venv/bin/python3
"""
将 Markdown 报告转为专业排版 Word 文档并上传飞书

用法: python3 md_to_feishu_docx.py <md文件路径> [飞书聊天ID]

功能特性:
  1. 使用 md2word 引擎，支持中文排版、专业样式
  2. 自动添加目录、页眉页脚
  3. 深蓝主题配色，表格表头蓝底白字，交替行浅灰
  4. 中文字体默认微软雅黑，代码块使用 Consolas
  5. 引用块蓝边+浅蓝底

需要环境变量: FEISHU_APP_ID, FEISHU_APP_SECRET
"""
import sys, os, json, re

MD_PATH = sys.argv[1] if len(sys.argv) > 1 else ""
CHAT_ID = sys.argv[2] if len(sys.argv) > 2 else os.getenv("FEISHU_CHAT_ID", "oc_99961a56e530e89f7e369cd6ecb50218")

if not MD_PATH or not os.path.exists(MD_PATH):
    print(f"❌ 文件不存在: {MD_PATH}")
    sys.exit(1)

# ── 1. 配置路径 ──
DOCX_PATH = MD_PATH.replace(".md", ".docx")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(BASE_DIR, "md2word_config.json")

# ── 2. 用 md2word 转换 ──
from md2word import convert_file, Config

# 加载配置
config = None
if os.path.exists(CONFIG_PATH):
    try:
        config = Config.from_file(CONFIG_PATH)
        print(f"📋 已加载排版配置: {CONFIG_PATH}")
    except Exception as e:
        print(f"⚠️ 配置加载失败，使用默认配置: {e}")

# 执行转换
try:
    if config:
        convert_file(MD_PATH, DOCX_PATH, config=config, toc=True)
    else:
        convert_file(MD_PATH, DOCX_PATH, toc=True)
    print(f"✅ Word 已生成: {DOCX_PATH}")
except Exception as e:
    print(f"❌ md2word 转换失败: {e}")
    print("⚠️ 降级使用 python-docx 基础转换...")
    
    # 降级方案：简单的 python-docx 转换
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35
    
    # 读取并转换
    with open(MD_PATH, 'r') as f:
        content = f.read()
    
    for line in content.split('\n'):
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            for run in h.runs:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('|') and line.endswith('|'):
            pass  # 简单跳过表格
        elif line.strip():
            p = doc.add_paragraph(line.strip())
    
    doc.save(DOCX_PATH)
    print(f"✅ Word 已生成(降级): {DOCX_PATH}")

# ── 3. 上传飞书 ──
import httpx

FEISHU_APP_ID = os.getenv("FEISHU_APP_ID", "")
FEISHU_APP_SECRET = os.getenv("FEISHU_APP_SECRET", "")

if not FEISHU_APP_ID or not FEISHU_APP_SECRET:
    print("❌ 缺少 FEISHU_APP_ID 或 FEISHU_APP_SECRET")
    sys.exit(1)

# 3.1 获取 Token
resp = httpx.post(
    "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
    json={"app_id": FEISHU_APP_ID, "app_secret": FEISHU_APP_SECRET},
    timeout=10
)
token = resp.json().get("tenant_access_token", "")
if not token:
    print(f"❌ 获取 Token 失败: {resp.text}")
    sys.exit(1)

# 3.2 上传文件
headers = {"Authorization": f"Bearer {token}"}
file_name = os.path.basename(DOCX_PATH)

with open(DOCX_PATH, 'rb') as f:
    upload_resp = httpx.post(
        "https://open.feishu.cn/open-apis/im/v1/files",
        headers=headers,
        files={"file": (file_name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"file_type": "doc", "file_name": file_name},
        timeout=30
    )

upload_data = upload_resp.json()
if upload_data.get("code") != 0:
    print(f"❌ 上传失败: {upload_data}")
    sys.exit(1)

file_key = upload_data["data"]["file_key"]

# 3.3 发送文件消息
msg_resp = httpx.post(
    f"https://open.feishu.cn/open-apis/im/v1/messages?receive_id_type=chat_id",
    headers={**headers, "Content-Type": "application/json"},
    json={
        "receive_id": CHAT_ID,
        "msg_type": "file",
        "content": json.dumps({"file_key": file_key})
    },
    timeout=15
)

msg_data = msg_resp.json()
if msg_data.get("code") == 0:
    print(f"✅ Word 文档已发送到飞书 ({file_name})")
else:
    print(f"❌ 发送失败: {msg_data}")
    sys.exit(1)
